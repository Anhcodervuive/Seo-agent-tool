from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).with_name("SEO_Copilot_Implementation_Report_and_Test_Guide.docx")

NAVY = "17365D"
BLUE = "2E74B5"
LIGHT_BLUE = "EAF2F8"
PALE_BLUE = "F4F8FC"
GREEN = "E2F0D9"
AMBER = "FFF2CC"
RED = "FCE4D6"
GRAY = "5B6573"
WHITE = "FFFFFF"
CONTENT_WIDTH = 9360


def shade(cell, color):
    props = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    props.append(shd)


def cell_margin(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text, bold=False, color="000000", size=10.2):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell_margin(cell)


def set_table_widths(table, widths):
    table.autofit = False
    table_pr = table._tbl.tblPr
    tbl_w = table_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        table_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = table_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        table_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for grid_col, width in zip(grid.gridCol_lst, widths):
        grid_col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def add_rule(paragraph, color="B9CDE5"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "10")
    bottom.set(qn("w:space"), "8")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    p_pr.append(borders)


def add_para(doc, text="", style=None, bold_lead=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(6)
    if bold_lead and text.startswith(bold_lead):
        run = p.add_run(bold_lead)
        run.bold = True
        p.add_run(text[len(bold_lead):])
    else:
        p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_callout(doc, title, text, color=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_widths(table, [CONTENT_WIDTH])
    cell = table.cell(0, 0)
    shade(cell, color)
    cell.text = ""
    cell_margin(cell, 150, 180, 150, 180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(title)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(NAVY)
    p2 = cell.add_paragraph(text)
    p2.paragraph_format.space_after = Pt(0)
    return table


def add_status_table(doc):
    rows = [
        ("17", "Crawl and ranking progress", "Completed", "Live progress shows crawled URLs, pending URLs, completed rank checks, and pending rank checks."),
        ("18", "Actual Meta Title in affected URLs", "Completed", "Shows the captured page title instead of crawler block-page text."),
        ("19", "Clearer meta description and content issue details", "Completed", "Shows relevant measurement fields, such as description length or word count, plus page content context."),
        ("20", "Duplicate H1 detail", "Completed", "Uses an H1 column and shows the captured H1 for each affected URL."),
        ("21", "Actual Meta Title in Meta Tags Report", "Completed", "Displays captured meta titles and does not treat block-page text as a valid title."),
        ("22", "Accurate GA4 date filtering", "Completed", "Fetches GA4 data for the exact selected period and saves it to the snapshot."),
        ("23", "Accurate GSC date filtering", "Completed", "Fetches Search Console data for the exact period and selected report view, then saves it."),
        ("24", "IST time with AM/PM", "Completed", "Snapshot and report timestamps use IST and an AM/PM display format."),
        ("25", "Backlink profile detail", "Completed", "Adds backlink overview and detailed Backlinks, Referring Domains, and Anchor Text sections."),
        ("Week 2", "Independent scheduled ranking checks", "Completed", "Adds a standalone recurring rank-check schedule alongside scheduled full audits."),
    ]
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    set_table_widths(table, [750, 2450, 1150, 5010])
    headers = ["Reference", "Item", "Status", "Delivered result"]
    for cell, value in zip(table.rows[0].cells, headers):
        shade(cell, NAVY)
        set_cell_text(cell, value, bold=True, color=WHITE)
    set_repeat_table_header(table.rows[0])
    for reference, item, status, result in rows:
        cells = table.add_row().cells
        fill = GREEN if status == "Completed" else AMBER
        for cell in cells:
            shade(cell, "FFFFFF")
        shade(cells[2], fill)
        set_cell_text(cells[0], reference, bold=True, color=NAVY)
        set_cell_text(cells[1], item, bold=True)
        set_cell_text(cells[2], status, bold=True, color="375623")
        set_cell_text(cells[3], result)
    return table


def add_flow_table(doc):
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_widths(table, [2100, 4200, 3060])
    for cell, value in zip(table.rows[0].cells, ["Stage", "Scheduled Full Audit", "Scheduled Ranking Check"]):
        shade(cell, NAVY)
        set_cell_text(cell, value, bold=True, color=WHITE)
    set_repeat_table_header(table.rows[0])
    rows = [
        ("Schedule becomes due", "Creates a durable queued job and a new snapshot.", "Creates a durable queued job and a new snapshot."),
        ("Website crawl", "Runs the configured crawl mode and captures crawl pages, links, images, issues, and structured data.", "Not run."),
        ("Google analytics data", "Collects GA4 and GSC data for the project configuration.", "Not run."),
        ("Keyword positions", "Checks each tracked keyword for the project and configured competitors through DataForSEO.", "Checks each tracked keyword for the project and configured competitors through DataForSEO."),
        ("Backlink and competitor metrics", "Collects the project backlink profile and competitor insight data.", "Not run."),
        ("Report", "Generates the analysis report after collection finishes.", "No report is generated; the snapshot records ranking-only results."),
    ]
    for label, full, rank in rows:
        cells = table.add_row().cells
        shade(cells[0], LIGHT_BLUE)
        set_cell_text(cells[0], label, bold=True, color=NAVY)
        set_cell_text(cells[1], full)
        set_cell_text(cells[2], rank)
    return table


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string("1F2937")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12

    for style_name, size, color, before, after in [
        ("Title", 25, NAVY, 0, 4),
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, NAVY, 12, 6),
        ("Heading 3", 11.5, NAVY, 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = style_name != "Title"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header.add_run("SEO Copilot | Implementation Report and Test Guide")
    header_run.font.size = Pt(8.5)
    header_run.font.color.rgb = RGBColor.from_string(GRAY)
    add_rule(header, "D9E2F3")
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("Internal implementation summary | August 2026")
    footer_run.font.size = Pt(8.5)
    footer_run.font.color.rgb = RGBColor.from_string(GRAY)


def main():
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(18)
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("SEO Copilot")
    run.font.name = "Calibri"
    run.font.size = Pt(28)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(NAVY)
    subtitle = doc.add_paragraph("Implementation Report and Test Guide")
    subtitle.paragraph_format.space_after = Pt(12)
    subtitle.runs[0].font.size = Pt(16)
    subtitle.runs[0].font.color.rgb = RGBColor.from_string(BLUE)
    meta = doc.add_paragraph("Scope: Requirements 17-25, GA4/GSC data accuracy, backlink reporting, and automated audit schedules")
    meta.paragraph_format.space_after = Pt(18)
    meta.runs[0].font.color.rgb = RGBColor.from_string(GRAY)
    add_rule(meta)

    add_callout(
        doc,
        "Implementation summary",
        "The completed work improves data accuracy, makes technical issue evidence easier to inspect, adds detailed backlink reporting, and provides separate recurring schedules for full audits and ranking-only checks.",
    )

    doc.add_heading("1. Executive Status", level=1)
    add_para(doc, "The implementation items in scope are complete and ready for end-to-end testing in the deployment environment. The main operational dependency is that the application worker, LibreCrawl service, Google credentials, and DataForSEO credentials are all available and healthy.")
    add_status_table(doc)

    doc.add_heading("2. GA4 Data Discrepancy Investigation", level=1)
    add_para(doc, "The original GA4 discrepancy was caused by date filtering being applied only in the interface. GA4 data was collected as one aggregated result when a snapshot was created. When a user selected another date range later, the screen changed the displayed dates but continued showing the old snapshot result.")
    add_para(doc, "For example, metrics collected for 6 July-3 August could appear while the interface displayed 1 July-9 August. The GA4 dashboard returned different values for the selected period, which confirmed that the Google property and credentials were working correctly.")
    add_callout(doc, "Implemented solution", "When a user applies a GA4 date range, the platform requests fresh GA4 data for that exact period, saves the returned rows to the selected snapshot, and then displays those stored rows. The screen now also gives loading, success, timeout, and error feedback.", GREEN)
    doc.add_heading("Expected result", level=2)
    add_bullets(doc, [
        "Total Users and Sessions should match the GA4 interface when the same property, date range, and Session primary channel group dimension are used.",
        "Metrics may still differ if the GA4 interface uses a different reporting identity, comparison, filter, timezone, or a different channel dimension.",
        "Each project should have an explicit Google account and GA4 property configuration before production testing.",
    ])

    doc.add_heading("3. GSC Data Accuracy", level=1)
    add_para(doc, "The same pattern applied to Search Console. The platform now requests fresh GSC data for the exact start date, end date, and selected report view (for example Queries or URLs), then saves the returned rows to the snapshot for later reuse.")
    add_bullets(doc, [
        "Long queries and long URLs are handled safely when saved, preventing the previous filtering error.",
        "The user receives loading, success, timeout, and error feedback instead of waiting without context.",
        "For a valid comparison, use the same Search Console property, search type, date range, and view in both products.",
    ])

    doc.add_heading("4. Requirement Delivery Detail", level=1)
    requirements = [
        ("Requirement 17 - Show Crawl and Ranking Progress", [
            "The live analysis panel shows crawled URLs, pending URLs, completed ranking checks, and pending ranking checks.",
            "Audit History separates Crawled URLs from Crawl Issues so the two values are not confused.",
        ]),
        ("Requirement 18 - Actual Meta Title in Affected URLs", [
            "Affected URL details for Meta Title issues show the captured title of the respective page.",
            "Crawler access-limit or block-page text is no longer presented as a valid page title.",
        ]),
        ("Requirement 19 - Clear Measurement Context", [
            "Meta Description Over 200 Characters shows Meta Description Length and the full Meta Description in separate fields.",
            "Comparable issue types, including Low Word Count Pages, show the relevant measurement and readable context.",
        ]),
        ("Requirement 20 - Duplicate H1 Evidence", [
            "The affected URL view uses an H1 column for Duplicate H1 issues.",
            "Each row shows the H1 captured for that URL rather than a generic Details field.",
        ]),
        ("Requirement 21 - Meta Tags Report Title Quality", [
            "The Meta Tags Report displays the captured Meta Title for each page.",
            "Block-page text is excluded from valid Meta Title values.",
        ]),
        ("Requirement 22 - Fresh GA4 Data", [
            "GA4 is retrieved for the exact date range selected by the user and saved to the snapshot.",
            "The UI includes loading, success, timeout, and error feedback.",
        ]),
        ("Requirement 23 - Fresh GSC Data", [
            "GSC is retrieved for the exact selected date range and report view and saved to the snapshot.",
            "Long data values no longer break the retrieval and storage flow.",
        ]),
        ("Requirement 24 - IST Timestamps", [
            "Snapshot and report timestamps display in Indian Standard Time using AM/PM format.",
            "The format is used consistently in Audit History, Snapshot Detail, and report-related pages.",
        ]),
        ("Requirement 25 - Detailed Backlink Profile", [
            "The Backlinks area provides summary metrics for total backlinks, referring domains, new backlinks, and lost backlinks.",
            "Detailed tabs provide Backlinks, Referring Domains, and Anchor Text evidence, including fields such as domain rank, target URL, link type, and first/last seen dates where available from DataForSEO.",
        ]),
    ]
    for heading, bullets in requirements:
        doc.add_heading(heading, level=2)
        add_bullets(doc, bullets)

    doc.add_heading("5. Automated Schedules", level=1)
    add_para(doc, "Project Settings now has a dedicated Schedules step. Scheduled Full Audit and Scheduled Ranking Check are configured independently and can each run Daily, Weekly, or Monthly. The application uses IST for schedule display.")
    add_flow_table(doc)
    add_callout(doc, "Queue safety", "When a schedule becomes due, the worker creates a durable job and snapshot in PostgreSQL. A project cannot have more than one queued or running job at the same time. If a full audit is already running, a due ranking check waits for a later worker cycle instead of creating duplicate API calls or duplicate data.", AMBER)

    doc.add_heading("6. Deployment Preparation", level=1)
    add_para(doc, "Before enabling schedules on the server, deploy the latest code, apply database migrations, and confirm that all background services are running. The worker is required; without it, scheduled jobs remain queued and do not execute.")
    doc.add_heading("Deployment steps", level=2)
    add_numbered(doc, [
        "Pull the latest code on the server.",
        "Build and restart the web and worker containers.",
        "Apply the database migrations.",
        "Confirm web, worker, and database containers are running.",
        "Confirm LibreCrawl is running and reachable from the worker container.",
        "Verify DataForSEO credit and Google permissions before enabling recurring work.",
    ])
    code_table = doc.add_table(rows=1, cols=1)
    set_table_widths(code_table, [CONTENT_WIDTH])
    cell = code_table.cell(0, 0)
    shade(cell, "1E293B")
    commands = "cd /opt/seo-agent-test\n" \
               "git pull --ff-only\n\n" \
               "cd pipeline\n" \
               "docker compose build web worker\n" \
               "docker compose up -d --force-recreate web worker\n" \
               "docker compose exec -T web python -m flask --app manage.py db upgrade\n\n" \
               "docker compose ps\n" \
               "docker compose logs -f worker"
    set_cell_text(cell, commands, color=WHITE, size=9.2)
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Consolas"

    doc.add_heading("Required production checks", level=2)
    add_bullets(doc, [
        "The worker container is Up and its log includes an audit-worker startup message.",
        "LibreCrawl is Up and LIBRECRAWL_URL is reachable from the worker container.",
        "DATAFORSEO_LOGIN and DATAFORSEO_PASSWORD are present and the DataForSEO account has credit.",
        "The configured GA4 credential has read access to the selected GA4 property.",
        "The configured GSC credential has read access to the exact Search Console property string.",
        "At least one tracked keyword exists before enabling Scheduled Ranking Check.",
    ])

    doc.add_heading("7. Test Guide", level=1)
    doc.add_heading("A. Validate GA4", level=2)
    add_numbered(doc, [
        "Open a completed snapshot and go to GA4 Metrics.",
        "Select a date range that differs from the original snapshot period and select Channel view.",
        "Click Apply. Confirm the loading message appears, then changes to success or a clear error message.",
        "Compare Total Users and Sessions with GA4 using the same property, dates, and Session primary channel group dimension.",
        "Refresh the page and confirm the retrieved result remains available for that snapshot.",
    ])
    doc.add_heading("B. Validate GSC", level=2)
    add_numbered(doc, [
        "Open GSC Queries in a completed snapshot.",
        "Select a new date range and choose Queries or URLs.",
        "Click Apply and wait for the loading and success message.",
        "Compare clicks and impressions against Search Console using the same property, dates, Search type, and selected view.",
        "Test a URL or query view containing long values to confirm the request completes without a storage error.",
    ])
    doc.add_heading("C. Validate issue evidence", level=2)
    add_numbered(doc, [
        "Open Meta Title duplicate and confirm the affected URL table shows the actual Meta Title.",
        "Open Meta Description Over 200 Characters and confirm the table shows a Meta Description Length column plus the full description.",
        "Open Duplicate H1 and confirm the table label is H1 and each row shows the captured H1.",
        "Open Meta Tags Report and confirm actual titles are shown rather than crawler block-page text.",
    ])
    doc.add_heading("D. Validate backlinks", level=2)
    add_numbered(doc, [
        "Run a new full audit after DataForSEO credentials and credit are confirmed.",
        "Open Snapshot Detail and locate the Backlinks section.",
        "Confirm total backlinks, referring domains, new in period, and lost in period display a value.",
        "Open Backlinks, Referring Domains, and Anchor Text tabs. Confirm details are readable and fields populate where DataForSEO supplies them.",
    ])
    doc.add_heading("E. Validate schedules", level=2)
    add_numbered(doc, [
        "Open Project Settings, select Schedules, enable one schedule, choose a frequency, and save changes.",
        "Confirm the next scheduled time appears in IST after reloading Settings.",
        "Watch docker compose logs -f worker. When the schedule is due, confirm the worker logs queue maintenance and job start messages.",
        "For Scheduled Full Audit, confirm a new snapshot includes crawl, GA4, GSC, rankings, backlinks, and a report result.",
        "For Scheduled Ranking Check, confirm a new snapshot refreshes ranking results only and does not run crawl, GA4, GSC, backlinks, or report generation.",
        "Enable both schedules if needed. Confirm only one job runs for the project at a time.",
    ])

    doc.add_heading("8. Operational Notes", level=1)
    add_bullets(doc, [
        "Older snapshots retain the data quality and crawler response that existed when they were created. Create a new snapshot after deployment for final validation.",
        "A partial status means one or more collection stages reported an error; use the Snapshot Notes and worker log to identify the specific service.",
        "Ranking, backlink, and competitor data depend on DataForSEO response availability and account credit.",
        "GA4 and GSC require correctly assigned project credentials and permissions. Crawl success does not grant Google API access.",
    ])
    add_callout(doc, "Recommended rollout", "Deploy first, complete one manual full audit for each project, validate GA4/GSC/Backlinks, and then enable recurring schedules. Start with Weekly ranking checks and Weekly or Monthly full audits to control API usage.", PALE_BLUE)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
