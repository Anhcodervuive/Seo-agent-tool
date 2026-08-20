"""Build a polished SEO Copilot client guide DOCX from a Markdown source.

The Markdown guide remains the source of truth. This builder turns its headings,
lists, callouts, and real data tables into a client-facing Word handoff. Run it
with the bundled Codex Python runtime documented in docs/user-guides/README.md.
"""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
DEFAULT_SOURCE = ROOT / "docs" / "user-guides" / "seo-copilot-client-guide-vi.md"
DEFAULT_OUTPUT = ROOT / "docs" / "user-guides" / "SEO_Copilot_Client_User_Guide_VI.docx"

PAGE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
COLORS = {
    "navy": "0B2545",
    "blue": "2E74B5",
    "dark_blue": "1F4D78",
    "muted": "5B6573",
    "line": "D6DEE8",
    "header_fill": "E8EEF5",
    "callout_fill": "F4F6F9",
    "cover_fill": "F7FAFE",
    "white": "FFFFFF",
}

COPY = {
    "vi": {
        "header": "HƯỚNG DẪN SỬ DỤNG",
        "footer": "SEO Copilot • Client User Guide • Trang ",
        "toc": "Mục lục",
        "toc_marker": "## Mục lục",
        "subject": "Hướng dẫn sử dụng Dashboard và AI SEO Copilot",
        "fallback_metadata": [
            ("Phiên bản", "1.0"),
            ("Cập nhật", "20 tháng 08 năm 2026"),
            ("Dành cho", "Client và team vận hành"),
            ("Phạm vi", "Week 3 Dashboard và AI Copilot"),
        ],
        "cover_callout_label": "Cách dùng tài liệu:",
        "cover_callout": "Audit thu thập dữ liệu; Dashboard cho bạn đọc dữ liệu; AI Copilot giúp giải thích dữ liệu đó. Hãy dùng hướng dẫn này như một checklist từ lần chạy đầu tiên đến việc ra quyết định hàng tháng.",
        "flow_steps": (
            ("1", "Run Analysis", "Thu thập dữ liệu"),
            ("2", "Snapshot", "Lưu mốc audit"),
            ("3", "Dashboard", "Xem thay đổi"),
            ("4", "AI Copilot", "Hỏi và ưu tiên việc"),
        ),
    },
    "en": {
        "header": "USER GUIDE",
        "footer": "SEO Copilot • Client User Guide • Page ",
        "toc": "Table of Contents",
        "toc_marker": "## Table of Contents",
        "subject": "SEO Copilot Dashboard and AI Client User Guide",
        "fallback_metadata": [
            ("Version", "1.0"),
            ("Updated", "20 August 2026"),
            ("Audience", "Clients and operations teams"),
            ("Scope", "Week 3 Dashboard and AI Copilot"),
        ],
        "cover_callout_label": "How to use this guide:",
        "cover_callout": "An audit collects data; the dashboard helps you read it; and AI Copilot explains it. Use this guide as a checklist from the first audit through monthly decision-making.",
        "flow_steps": (
            ("1", "Run Analysis", "Collect the data"),
            ("2", "Snapshot", "Save an audit checkpoint"),
            ("3", "Dashboard", "Review the changes"),
            ("4", "AI Copilot", "Ask and prioritize work"),
        ),
    },
}


def set_run_font(run, *, size=None, color=None, bold=None, italic=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=COLORS["line"], size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa, *, indent_dxa=TABLE_INDENT_DXA):
    """Set fixed Word geometry: tblW, tblInd, tblGrid and every cell width."""
    if sum(widths_dxa) != PAGE_WIDTH_DXA:
        raise ValueError(f"Table widths must sum to {PAGE_WIDTH_DXA}, got {sum(widths_dxa)}")
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(PAGE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for node in list(grid):
        grid.remove(node)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)


def mark_first_row_as_header(table):
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        tr_pr.append(OxmlElement("w:tblHeader"))


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, separate, text, end])


def set_cell_text(cell, text, *, bold=False, color=COLORS["navy"], size=10, align=WD_ALIGN_PARAGRAPH.LEFT):
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = Pt(12.5)
    paragraph.clear()
    add_inline_runs(paragraph, text, size=size, color=color, force_bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell)


INLINE_TOKEN = re.compile(r"(\*\*.+?\*\*|`.+?`)")


def add_inline_runs(paragraph, text, *, size=11, color=COLORS["navy"], force_bold=False, italic=False):
    for piece in INLINE_TOKEN.split(text):
        if not piece:
            continue
        bold = force_bold
        is_code = piece.startswith("`") and piece.endswith("`")
        if piece.startswith("**") and piece.endswith("**"):
            piece, bold = piece[2:-2], True
        elif is_code:
            piece = piece[1:-1]
        run = paragraph.add_run(piece)
        set_run_font(run, size=size, color=color, bold=bold, italic=italic)
        if is_code:
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
            run._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")


def set_paragraph_body(paragraph, *, after=6, before=0, line=15, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    paragraph.alignment = alignment
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = Pt(line)
    paragraph.paragraph_format.widow_control = True


def set_document_styles(document):
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(COLORS["navy"])
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = Pt(15)

    style_specs = {
        "Heading 1": (16, COLORS["blue"], 18, 10),
        "Heading 2": (13, COLORS["blue"], 14, 7),
        "Heading 3": (12, COLORS["dark_blue"], 10, 5),
    }
    for name, (size, color, before, after) in style_specs.items():
        style = document.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.widow_control = True


def configure_section(section, copy):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    left = paragraph.add_run("SEO COPILOT")
    set_run_font(left, size=8.5, color=COLORS["blue"], bold=True)
    spacer = paragraph.add_run(f"  |  {copy['header']}")
    set_run_font(spacer, size=8.5, color=COLORS["muted"], bold=True)

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    label = paragraph.add_run(copy["footer"])
    set_run_font(label, size=8.5, color=COLORS["muted"])
    add_page_field(paragraph)


def add_callout(document, text, *, label=None):
    table = document.add_table(rows=1, cols=1)
    # Named override: callouts use extra 180 DXA inner start padding, so their
    # table indent matches the visible body-text alignment of the callout.
    set_table_geometry(table, [PAGE_WIDTH_DXA], indent_dxa=180)
    set_table_borders(table, color="B7CBE0", size="8")
    mark_first_row_as_header(table)
    cell = table.cell(0, 0)
    shade_cell(cell, COLORS["callout_fill"])
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = Pt(14)
    if label:
        run = paragraph.add_run(f"{label} ")
        set_run_font(run, size=10.5, color=COLORS["dark_blue"], bold=True)
    add_inline_runs(paragraph, text, size=10.5, color=COLORS["navy"])
    set_cell_margins(cell, top=110, bottom=110, start=180, end=180)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_data_table(document, rows):
    if not rows:
        return
    column_count = len(rows[0])
    if column_count == 2:
        widths = [2700, 6660]
    elif column_count == 3:
        widths = [2200, 3000, 4160]
    elif column_count == 4:
        widths = [1800, 2520, 2520, 2520]
    else:
        widths = [PAGE_WIDTH_DXA // column_count] * column_count
        widths[-1] += PAGE_WIDTH_DXA - sum(widths)
    table = document.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    set_table_borders(table)
    mark_first_row_as_header(table)

    for row_index, values in enumerate(rows):
        for cell, value in zip(table.rows[row_index].cells, values):
            is_header = row_index == 0
            if is_header:
                shade_cell(cell, COLORS["header_fill"])
            set_cell_text(
                cell,
                value,
                bold=is_header,
                color=COLORS["dark_blue"] if is_header else COLORS["navy"],
                size=9.5,
                align=WD_ALIGN_PARAGRAPH.CENTER if is_header else WD_ALIGN_PARAGRAPH.LEFT,
            )
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def parse_table(lines, start):
    rows = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        values = [cell.strip() for cell in lines[index].strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in values):
            rows.append(values)
        index += 1
    return rows, index


def collect_sections(lines, toc_marker):
    return [line[3:].strip() for line in lines if line.startswith("## ") and not line.startswith(toc_marker)]


def add_cover(document, title, subtitle, metadata, sections, copy):
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_before = Pt(32)
    spacer.paragraph_format.space_after = Pt(14)
    kicker = document.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.LEFT
    kicker.paragraph_format.space_after = Pt(4)
    run = kicker.add_run("CUSTOMER ENABLEMENT GUIDE")
    set_run_font(run, size=10, color=COLORS["blue"], bold=True)

    title_p = document.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(8)
    title_p.paragraph_format.line_spacing = Pt(34)
    run = title_p.add_run(title)
    set_run_font(run, size=30, color=COLORS["navy"], bold=True)

    subtitle_p = document.add_paragraph()
    subtitle_p.paragraph_format.space_before = Pt(0)
    subtitle_p.paragraph_format.space_after = Pt(20)
    subtitle_p.paragraph_format.line_spacing = Pt(20)
    add_inline_runs(subtitle_p, subtitle, size=13, color=COLORS["muted"])

    table = document.add_table(rows=2, cols=2)
    # Named override: the cover metadata grid has the same 180 DXA padding.
    set_table_geometry(table, [4680, 4680], indent_dxa=180)
    set_table_borders(table, color="D7E2EF", size="8")
    mark_first_row_as_header(table)
    for index, (label, value) in enumerate(metadata):
        row, col = divmod(index, 2)
        cell = table.cell(row, col)
        shade_cell(cell, COLORS["cover_fill"])
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(1)
        label_run = paragraph.add_run(label.upper())
        set_run_font(label_run, size=8.5, color=COLORS["blue"], bold=True)
        value_p = cell.add_paragraph()
        value_p.paragraph_format.space_before = Pt(0)
        value_p.paragraph_format.space_after = Pt(0)
        add_inline_runs(value_p, value, size=10.5, color=COLORS["navy"], force_bold=True)
        set_cell_margins(cell, top=140, bottom=140, start=180, end=180)

    document.add_paragraph().paragraph_format.space_after = Pt(6)
    add_callout(document, copy["cover_callout"], label=copy["cover_callout_label"])

    flow = document.add_table(rows=1, cols=4)
    # Named override: the compact four-step process strip uses 100 DXA cells.
    set_table_geometry(flow, [2340, 2340, 2340, 2340], indent_dxa=100)
    set_table_borders(flow, color="C6D4E3", size="6")
    mark_first_row_as_header(flow)
    for cell, (number, label, detail) in zip(flow.rows[0].cells, copy["flow_steps"]):
        shade_cell(cell, COLORS["cover_fill"])
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(2)
        run = paragraph.add_run(number)
        set_run_font(run, size=14, color=COLORS["blue"], bold=True)
        label_p = cell.add_paragraph()
        label_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        label_p.paragraph_format.space_after = Pt(1)
        add_inline_runs(label_p, label, size=9.5, color=COLORS["navy"], force_bold=True)
        detail_p = cell.add_paragraph()
        detail_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        detail_p.paragraph_format.space_after = Pt(0)
        add_inline_runs(detail_p, detail, size=8.5, color=COLORS["muted"])
        set_cell_margins(cell, top=100, bottom=100, start=100, end=100)

    document.add_page_break()
    toc_heading = document.add_paragraph(style="Heading 1")
    toc_heading.add_run(copy["toc"])
    toc_heading.paragraph_format.space_before = Pt(0)
    toc_heading.paragraph_format.space_after = Pt(10)
    for section in sections:
        paragraph = document.add_paragraph(style="List Number")
        set_paragraph_body(paragraph, after=4, line=14)
        paragraph.paragraph_format.left_indent = Inches(0.375)
        paragraph.paragraph_format.first_line_indent = Inches(-0.188)
        add_inline_runs(paragraph, section, size=10.5, color=COLORS["navy"])
    document.add_page_break()


def build_document(source=DEFAULT_SOURCE, output=DEFAULT_OUTPUT):
    source, output = Path(source), Path(output)
    lines = source.read_text(encoding="utf-8").splitlines()
    language = "en" if any(line.startswith(COPY["en"]["toc_marker"]) for line in lines) else "vi"
    copy = COPY[language]
    title = next((line[2:].strip() for line in lines if line.startswith("# ")), "SEO Copilot")
    subtitle = next((line[2:].strip() for line in lines if line.startswith("> ")), "")
    metadata = []
    for line in lines:
        match = re.match(r"\*\*(.+?):\*\*\s*(.+?)\s*$", line)
        if match:
            metadata.append((match.group(1), match.group(2).rstrip("  ")))
        if len(metadata) == 4:
            break
    metadata = metadata or copy["fallback_metadata"]

    document = Document()
    configure_section(document.sections[0], copy)
    set_document_styles(document)
    document.core_properties.title = title
    document.core_properties.subject = copy["subject"]
    document.core_properties.author = "SEO Copilot"
    document.core_properties.keywords = "SEO, dashboard, audit, AI Copilot, client guide"
    document.core_properties.comments = "Client-facing guide generated from Markdown source of truth."

    add_cover(document, title, subtitle, metadata, collect_sections(lines, copy["toc_marker"]), copy)

    toc_index = next((i for i, line in enumerate(lines) if line.startswith(copy["toc_marker"])), -1)
    index = next(
        (i for i, line in enumerate(lines[toc_index + 1 :], start=toc_index + 1) if line.startswith("## ")),
        len(lines),
    )
    while index < len(lines):
        raw = lines[index]
        stripped = raw.strip()
        if not stripped:
            index += 1
            continue
        if raw.startswith(copy["toc_marker"]):
            index += 1
            while index < len(lines) and not lines[index].startswith("## "):
                index += 1
            continue
        if raw.startswith("## "):
            paragraph = document.add_paragraph(style="Heading 1")
            add_inline_runs(paragraph, raw[3:].strip(), size=16, color=COLORS["blue"], force_bold=True)
            index += 1
            continue
        if raw.startswith("### "):
            paragraph = document.add_paragraph(style="Heading 2")
            add_inline_runs(paragraph, raw[4:].strip(), size=13, color=COLORS["blue"], force_bold=True)
            index += 1
            continue
        if raw.startswith("#### "):
            paragraph = document.add_paragraph(style="Heading 3")
            add_inline_runs(paragraph, raw[5:].strip(), size=12, color=COLORS["dark_blue"], force_bold=True)
            index += 1
            continue
        if stripped.startswith("|"):
            rows, index = parse_table(lines, index)
            add_data_table(document, rows)
            continue
        if stripped.startswith("> "):
            callout = stripped[2:].strip()
            match = re.match(r"\*\*(.+?):\*\*\s*(.*)", callout)
            if match:
                add_callout(document, match.group(2), label=f"{match.group(1)}:")
            else:
                add_callout(document, callout)
            index += 1
            continue
        bullet = re.match(r"^-\s+(.+)", stripped)
        numbered = re.match(r"^\d+\.\s+(.+)", stripped)
        if bullet or numbered:
            paragraph = document.add_paragraph(style="List Bullet" if bullet else "List Number")
            set_paragraph_body(paragraph, after=4, line=15)
            paragraph.paragraph_format.left_indent = Inches(0.375)
            paragraph.paragraph_format.first_line_indent = Inches(-0.188)
            add_inline_runs(paragraph, (bullet or numbered).group(1), size=11, color=COLORS["navy"])
            index += 1
            continue
        if stripped.startswith("---"):
            index += 1
            continue
        if stripped.startswith("**") and stripped.endswith("**"):
            paragraph = document.add_paragraph()
            set_paragraph_body(paragraph, after=6, line=15)
            add_inline_runs(paragraph, stripped, size=10.5, color=COLORS["muted"])
            index += 1
            continue
        if stripped:
            paragraph = document.add_paragraph()
            set_paragraph_body(paragraph)
            add_inline_runs(paragraph, stripped, size=11, color=COLORS["navy"])
        index += 1

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    return output


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--source", type=Path, default=DEFAULT_SOURCE)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    args = parser.parse_args()
    target = build_document(args.source, args.output)
    print(f"Created {target}")
